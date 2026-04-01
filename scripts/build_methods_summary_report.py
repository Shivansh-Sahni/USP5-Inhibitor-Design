from __future__ import annotations

from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "outputs"
FINAL_MODEL_DIR = OUTPUT_DIR / "final_model"
MD_PATH = OUTPUT_DIR / "methods_features_and_narrowing_report.md"
DOCX_PATH = OUTPUT_DIR / "methods_features_and_narrowing_report.docx"


REGRESSION_FEATURES = [
    ("mw", "Molecular weight", "Tracks overall size and strongly affects permeability, potency trends, and developability."),
    ("logp", "logP", "Captures lipophilicity, which often influences membrane penetration and nonspecific binding."),
    ("tpsa", "Topological polar surface area", "Measures polarity and is useful for absorption and binding-environment balance."),
    ("hbd", "Hydrogen-bond donors", "Counts donor groups that can drive binding or reduce permeability if too high."),
    ("hba", "Hydrogen-bond acceptors", "Counts acceptor groups that affect recognition and physicochemical behavior."),
    ("rot", "Rotatable bonds", "Approximates flexibility, which can affect entropy, oral exposure, and 3D fit."),
    ("rings", "Ring count", "Tracks scaffold rigidity and shape complexity."),
    ("hac", "Heavy atom count", "A simple size and complexity descriptor that often correlates with potency trends."),
    ("fsp3", "Fraction sp3 carbons", "Measures saturation / 3D character and can help distinguish flat from more three-dimensional molecules."),
    ("bertz", "Bertz complexity", "Summarizes structural complexity from the molecular graph."),
    ("balaban", "Balaban J index", "A graph-connectivity descriptor that reflects topology."),
    ("chi0v", "Valence connectivity index Chi0v", "Encodes atom connectivity at a local graph level."),
    ("chi1v", "Valence connectivity index Chi1v", "Encodes one-bond connectivity patterns."),
    ("chi2v", "Valence connectivity index Chi2v", "Encodes two-bond graph connectivity patterns."),
    ("kappa1", "Kier shape index Kappa1", "Represents coarse molecular shape."),
    ("kappa2", "Kier shape index Kappa2", "Represents shape and branching at a deeper level."),
    ("kappa3", "Kier shape index Kappa3", "Represents higher-order shape complexity and was the top single feature in the final fit."),
]

ENUMERATION_TECHNIQUES = [
    ("R-group enumeration", "Keep a core fragment fixed and swap side fragments around it.", "Implemented with BRICS core fragments plus a pooled side-fragment set, rebuilding products while preserving the parent core.", "1205"),
    ("Reaction-based enumeration", "Generate analogs only through specific chemistry transforms.", "Implemented with RDKit reaction SMARTS such as acid-to-amide conversions and phenol O-alkylation.", "31"),
    ("Reagent-pool combinatorics", "Combine a shared pool of compatible core and side fragments to make many plausible analogs.", "Implemented with a limited BRICS fragment pool and BRICSBuild to create global combinations.", "1200"),
    ("Matched molecular pair expansion", "Make small medicinal-chemistry edits such as halogen swaps or methoxy/hydroxyl interconversion.", "Implemented with targeted single-step SMARTS replacements like Cl to F, Br to Cl, and OMe to OH.", "25"),
    ("Bioisosteric swaps", "Replace one functional group with a chemically related surrogate.", "Implemented with transforms such as acid to amide, acid to thioacid, and amide to thioamide.", "15"),
    ("Linker scans", "Vary the connector between motifs to alter spacing, polarity, or flexibility.", "Implemented with predefined SMARTS transforms that insert oxygen or carbamate-like changes into existing linkers.", "5"),
    ("Ring-size / ring-system scans", "Change ring size or ring composition to test shape and polarity changes.", "Implemented with SMARTS transforms such as pyrrolidine to piperidine and piperidine to morpholine.", "3"),
    ("Heteroatom walks", "Move from aryl CH positions into heteroaryl nitrogens to tune polarity and recognition.", "Implemented as aromatic CH to N conversions, capped per parent to avoid runaway expansion.", "61"),
    ("Scaffold hopping", "Preserve side chains but replace the central core with an alternative BRICS core.", "Implemented by pairing parent-derived side fragments with non-native core fragments from other positives.", "671"),
    ("Fragment growing", "Start from a core fragment and add one side fragment outward.", "Implemented with limited BRICS builds of one core plus one side fragment to create smaller controlled expansions.", "430"),
]

NARROWING_TECHNIQUES = [
    ("Model-first potency screen", "The saved ExtraTrees regression model scores the full library first, and only compounds with predicted pIC50 >= 4.60 and acceptable descriptor-space distance continue."),
    ("PAINS and BRENK alert removal", "Compounds with nuisance or problematic substructure alerts are removed before deeper triage."),
    ("Lipinski-style developability filter", "The pipeline keeps compounds with <= 1 Lipinski violation to avoid drifting too far into poor oral-drug-like space."),
    ("Topological polar surface area filter", "TPSA is constrained to 40-140 A^2 to balance permeability with needed polarity."),
    ("Molecular surface area filter", "Labute approximate surface area is constrained to 130-235 A^2 to avoid structures that are too small or too bulky for the intended pocket and exposure profile."),
    ("Molecular weight filter", "MW is constrained to 250-550 to remove very small fragments and oversized analogs."),
    ("Flexibility and charge filter", "Rotatable bonds are capped at 10 and absolute formal charge at 1 to keep molecules closer to plausible lead space."),
    ("ADMET-AI multigate", "The surviving set is screened with ADMET-AI for AMES, hERG, ClinTox, HIA, oral bioavailability, Caco-2, and solubility-related properties."),
    ("Multistructure 3D binding plausibility", "Compounds are compared against several USP5 ZnF-UBD co-crystal-inspired templates to assess binding-score, shape overlap, pharmacophore match, and steric clashes."),
    ("Portfolio selection", "The last step separates strict primary leads from orthogonal backups so the project does not overcommit to only one chemotype."),
]


def set_doc_defaults(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(20)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(12)


def add_table(doc: Document, dataframe: pd.DataFrame) -> None:
    table = doc.add_table(rows=1, cols=len(dataframe.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, column in enumerate(dataframe.columns):
        hdr[idx].text = str(column)
    for row in dataframe.itertuples(index=False):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)


def write_markdown() -> None:
    lead_counts = pd.read_csv(OUTPUT_DIR / "lead_selection_counts.csv")
    lines: list[str] = []
    lines.append("# Regression Features, Enumeration Methods, and Narrowing Pipeline")
    lines.append("")
    lines.append(f"Generated on {date.today().isoformat()}.")
    lines.append("")
    lines.append("## Project context")
    lines.append("")
    lines.append("- Final potency model: saved `ExtraTreesRegressor`")
    lines.append("- Final reported in-sample `R^2`: `0.893359`")
    lines.append("- Broad enumerated library: `3264` unique compounds")
    lines.append("- Final screening start set: `3274` compounds including original positives")
    lines.append("")
    lines.append("## 1. Features used for regression")
    lines.append("")
    lines.append("The final regression model used the `base_graph` feature block, meaning standard physicochemical descriptors plus graph-topology descriptors derived from RDKit.")
    lines.append("")
    for feature, name, explanation in REGRESSION_FEATURES:
        lines.append(f"- `{feature}` ({name}): {explanation}")
    lines.append("")
    lines.append("## 2. Enumeration techniques")
    lines.append("")
    lines.append("The canonical enumeration workflow used 10 methods.")
    lines.append("")
    for name, what_it_is, how_done, count in ENUMERATION_TECHNIQUES:
        lines.append(f"### {name}")
        lines.append("")
        lines.append(f"- What it does: {what_it_is}")
        lines.append(f"- How it was done in this project: {how_done}")
        lines.append(f"- Unique products produced in the final broad library: `{count}`")
        lines.append("")
    lines.append("## 3. Narrowing-down techniques")
    lines.append("")
    lines.append("The final narrowing funnel starts with potency prediction and then applies increasingly strict chemistry, ADMET, and 3D plausibility filters.")
    lines.append("")
    for name, description in NARROWING_TECHNIQUES:
        lines.append(f"### {name}")
        lines.append("")
        lines.append(description)
        lines.append("")
    lines.append("## 4. Final stage counts")
    lines.append("")
    for row in lead_counts.itertuples(index=False):
        lines.append(f"- `{row.stage}`: `{row.remaining_unique_compounds}`")
    lines.append("")
    lines.append("## 5. Verbal summary")
    lines.append("")
    lines.append("The regression model uses interpretable physicochemical and graph-shape descriptors, the enumeration layer spans 10 complementary medicinal-chemistry-inspired expansion methods, and the narrowing layer combines potency prediction, property filtering, ADMET-AI, and multistructure 3D binding plausibility. Together, these steps mean most of the core computational pipeline has already been completed.")
    lines.append("")
    MD_PATH.write_text("\n".join(lines))


def write_docx() -> None:
    lead_counts = pd.read_csv(OUTPUT_DIR / "lead_selection_counts.csv")
    enum_counts = pd.read_csv(OUTPUT_DIR / "enumeration_method_counts.csv")

    doc = Document()
    set_doc_defaults(doc)

    title = doc.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Regression Features, Enumeration Methods, and Narrowing Pipeline")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Compact project summary for mentor discussion").italic = True
    stamp = doc.add_paragraph()
    stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stamp.add_run(f"Generated on {date.today().isoformat()}")

    doc.add_paragraph(style="Heading 1").add_run("Project Context")
    for item in [
        "Final potency model: saved ExtraTreesRegressor",
        "Final reported in-sample R^2: 0.893359",
        "Broad enumerated library: 3264 unique compounds",
        "Final screening start set: 3274 compounds including original positives",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(style="Heading 1").add_run("1. Features Used for Regression")
    doc.add_paragraph("The final model used the base_graph feature set: standard physicochemical descriptors combined with graph-topology descriptors.")
    feature_df = pd.DataFrame(REGRESSION_FEATURES, columns=["feature", "name", "why it matters"])
    add_table(doc, feature_df)

    for image_name, caption in [
        ("predicted_vs_actual.png", "Final model predicted vs observed pIC50"),
        ("feature_importance_top12.png", "Top feature importances"),
    ]:
        path = FINAL_MODEL_DIR / image_name
        if path.exists():
            doc.add_picture(str(path), width=Inches(6.1))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.add_run(caption).italic = True

    doc.add_paragraph(style="Heading 1").add_run("2. Enumeration Techniques")
    doc.add_paragraph("The canonical enumeration workflow used 10 complementary methods. The table below explains each method and shows how many unique products it contributed to the final broad library.")
    enum_df = pd.DataFrame(ENUMERATION_TECHNIQUES, columns=["technique", "what it does", "how it was done", "unique products"])
    add_table(doc, enum_df)
    doc.add_paragraph("Final method counts from the canonical broad library:", style="Heading 2")
    add_table(doc, enum_counts.rename(columns={"method": "method", "unique_products": "count"}))

    doc.add_paragraph(style="Heading 1").add_run("3. Techniques Used to Narrow Down to Leads")
    doc.add_paragraph("The final narrowing funnel is intentionally ordered from fast high-throughput scoring to stricter developability and structural plausibility checks.")
    narrow_df = pd.DataFrame(NARROWING_TECHNIQUES, columns=["narrowing technique", "role in the pipeline"])
    add_table(doc, narrow_df)

    doc.add_paragraph(style="Heading 2").add_run("Final Funnel Counts")
    add_table(doc, lead_counts.rename(columns={"stage": "stage", "remaining_unique_compounds": "count"}))

    doc.add_paragraph(style="Heading 1").add_run("Short Summary")
    doc.add_paragraph(
        "The final workflow uses interpretable chemistry descriptors for regression, ten medicinal-chemistry-inspired expansion methods for enumeration, and a multistage narrowing pipeline built around potency, properties, ADMET-AI, and USP5 3D binding plausibility. This means the major computational steps of the project are already in place and ready to support mentor discussion, paper planning, and experimental prioritization."
    )

    doc.save(DOCX_PATH)


def main() -> None:
    write_markdown()
    write_docx()
    print(MD_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
