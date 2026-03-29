from __future__ import annotations

from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "outputs"
FINAL_MODEL_DIR = OUTPUT_DIR / "final_model"
REPORT_PATH = OUTPUT_DIR / "usp5_project_report_for_mentor.docx"


def add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    new_run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_doc_defaults(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(20)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(12)


def add_table(doc: Document, dataframe: pd.DataFrame, title: str | None = None) -> None:
    if title:
        doc.add_paragraph(title, style="Heading 2")
    table = doc.add_table(rows=1, cols=len(dataframe.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, column in enumerate(dataframe.columns):
        hdr[idx].text = str(column)
    for row in dataframe.itertuples(index=False):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if isinstance(value, float):
                text = f"{value:.4f}".rstrip("0").rstrip(".")
            else:
                text = str(value)
            cells[idx].text = text


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_report() -> None:
    pipeline_summary = (OUTPUT_DIR / "final_pipeline_summary.md").read_text()
    model_report = (FINAL_MODEL_DIR / "final_model_report.md").read_text()
    lead_summary = (OUTPUT_DIR / "lead_selection_summary.md").read_text()

    final_leads = pd.read_csv(OUTPUT_DIR / "final_leads.csv")
    backup_leads = pd.read_csv(OUTPUT_DIR / "backup_leads.csv")
    enum_counts = pd.read_csv(OUTPUT_DIR / "enumeration_method_counts.csv")
    lead_counts = pd.read_csv(OUTPUT_DIR / "lead_selection_counts.csv")
    modeling_df = pd.read_csv(OUTPUT_DIR / "modeling_dataset.csv")

    doc = Document()
    set_doc_defaults(doc)

    title = doc.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("USP5 Computational Project Report")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Mentor-facing summary of the modeling, enumeration, and lead selection workflow").italic = True
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(f"Generated on {date.today().isoformat()}")

    doc.add_paragraph()
    exec_summary = doc.add_paragraph(style="Heading 1")
    exec_summary.add_run("Executive Summary")
    add_bullets(
        doc,
        [
            "The project now has one canonical computational pipeline built around a saved ExtraTrees regression model, a broad 10-method enumeration library, and a multistage lead-selection funnel.",
            "The final saved potency model achieved in-sample R^2 = 0.893 on the current row-level USP5 dataset and remains the project-level potency engine.",
            "The broad enumeration workflow produced 3264 unique virtual compounds from the positive training chemistry space.",
            "The final screening pipeline narrowed 3274 total structures (broad enumeration plus original positives) to 4 strict primary leads and 5 orthogonal backup leads.",
            "Most computational cheminformatics work is complete; the main remaining tasks before paper writing are figure assembly, polished results framing, and ideally experimental validation planning.",
        ],
    )

    doc.add_paragraph(style="Heading 1").add_run("Project Aim")
    doc.add_paragraph(
        "The goal of this project is to build a practical small-data USP5 inhibitor discovery workflow that combines potency modeling, chemically meaningful virtual enumeration, and computational lead prioritization. The emphasis is on transparent and interpretable cheminformatics rather than black-box modeling."
    )

    doc.add_paragraph(style="Heading 1").add_run("Dataset Status")
    add_bullets(
        doc,
        [
            "Raw input file: data/raw/First.csv",
            f"Valid row-level entries used by the final model: 26",
            f"Deduplicated canonical molecules in the modeling table: {len(modeling_df)}",
            "The dataset mixes measured activity rows with assigned active/inactive labels; this is tracked explicitly in the cleaned outputs.",
        ],
    )
    add_table(
        doc,
        modeling_df[
            [
                "representative_id",
                "target_pIC50",
                "target_origin",
                "has_measured_row",
                "has_active_no_ic50_row",
                "has_inactive_row",
            ]
        ],
        title="Modeling Dataset Overview",
    )

    doc.add_paragraph(style="Heading 1").add_run("Final Potency Model")
    doc.add_paragraph(
        "The final potency model is the saved ExtraTreesRegressor reported in outputs/final_model/final_model_report.md. This model was deliberately preserved as the canonical project model for all final screening."
    )
    add_bullets(
        doc,
        [
            "Model family: ExtraTreesRegressor",
            "Feature block: physicochemical descriptors plus graph-topology descriptors",
            "Reported in-sample R^2: 0.893359",
            "Reported MAE: 0.176795",
            "Reported RMSE: 0.368224",
            "Model artifact: outputs/final_model/final_model.joblib",
        ],
    )
    for image_name, caption in [
        ("predicted_vs_actual.png", "Final model predicted vs observed pIC50"),
        ("feature_importance_top12.png", "Top feature importances in the final model"),
    ]:
        image_path = FINAL_MODEL_DIR / image_name
        if image_path.exists():
            doc.add_picture(str(image_path), width=Inches(6.3))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.add_run(caption).italic = True

    doc.add_paragraph(style="Heading 1").add_run("Enumeration Workflow")
    doc.add_paragraph(
        "The final enumeration layer is the broad 10-method chemistry workflow. This is the canonical virtual library generation stage used for downstream lead screening."
    )
    add_bullets(
        doc,
        [
            "Canonical enumeration script: scripts/run_enumeration_10_methods.py",
            "Canonical broad library: outputs/enumeration_library_10_methods.csv",
            "Total unique enumerated products: 3264",
        ],
    )
    add_table(doc, enum_counts.rename(columns={"unique_products": "count"}), title="Enumeration Method Breakdown")

    doc.add_paragraph(style="Heading 1").add_run("Final Lead Selection Workflow")
    doc.add_paragraph(
        "The canonical lead-selection workflow restores the saved final ExtraTrees model as the sole potency engine, then applies property filtering, ADMET-AI, and multistructure USP5 3D template-docking / pharmacophore plausibility."
    )
    add_bullets(
        doc,
        [
            "Canonical screening script: scripts/run_lead_selection.py",
            "Property filters include Lipinski-style limits, TPSA, molecular surface area, flexibility, and charge sanity.",
            "ADMET stage uses ADMET-AI with AMES, hERG, ClinTox, HIA, oral bioavailability, permeability, and solubility-related outputs.",
            "3D structural evidence uses USP5 ZnF-UBD co-crystal structures 6DXT, 7MS5, 7MS6, and 7MS7.",
            "The final output is split into primary leads and orthogonal backup leads to avoid overclaiming a single chemotype as the only viable program.",
        ],
    )
    add_table(doc, lead_counts.rename(columns={"remaining_unique_compounds": "count"}), title="Lead Selection Funnel")

    doc.add_paragraph(style="Heading 1").add_run("Primary Leads")
    primary_display = final_leads[
        [
            "product_smiles",
            "primary_parent_id",
            "pred_pIC50",
            "pred_ic50_uM",
            "AMES",
            "hERG",
            "best_binding_score",
            "scaffold",
        ]
    ].copy()
    add_table(doc, primary_display, title="Canonical Primary Lead Set")

    doc.add_paragraph(style="Heading 1").add_run("Orthogonal Backup Leads")
    backup_display = backup_leads[
        [
            "product_smiles",
            "primary_parent_id",
            "pred_pIC50",
            "AMES",
            "hERG",
            "best_binding_score",
            "scaffold",
        ]
    ].copy()
    add_table(doc, backup_display, title="Canonical Backup Lead Set")

    doc.add_paragraph(style="Heading 1").add_run("Interpretation of Current Results")
    add_bullets(
        doc,
        [
            "The strict primary lead evidence remains concentrated in the CHEMBL5278336 acid-sulfonamide family.",
            "The backup program retains orthogonal chemistry, especially CHEMBL5410606-derived bicyclic carbonyl analogs and a weaker heteroaryl-acid branch.",
            "The project now has a consistent folder structure and a single final pipeline rather than multiple competing screening variants.",
            "From a project-management standpoint, most of the computational chemistry work is now done.",
        ],
    )

    doc.add_paragraph(style="Heading 1").add_run("Limitations")
    add_bullets(
        doc,
        [
            "The final model is exploratory and small-data; strong prospective claims should be avoided.",
            "The final model performance reported here is in-sample rather than a large prospective benchmark.",
            "The 3D stage is a multistructure template-docking / pharmacophore surrogate, not a full production docking plus MD validation campaign.",
            "All lead calls remain computational priorities until experimentally tested.",
            "The dataset includes assigned labels in addition to measured values, which should be acknowledged in any presentation or manuscript.",
        ],
    )

    doc.add_paragraph(style="Heading 1").add_run("Work Completed")
    add_bullets(
        doc,
        [
            "Dataset parsing, cleaning, label annotation, and canonicalization",
            "Descriptor and fingerprint generation",
            "Baseline similarity and scaffold analysis",
            "Final exploratory USP5 potency model selection and report generation",
            "Broad 10-method combinatorial enumeration",
            "ADMET-AI integration",
            "Multistructure USP5 3D screening against experimentally determined co-crystal structures",
            "Primary and orthogonal backup lead prioritization",
            "Project-folder cleanup and canonical final pipeline consolidation",
        ],
    )

    doc.add_paragraph(style="Heading 1").add_run("What Remains Before Paper Writing")
    add_bullets(
        doc,
        [
            "Assemble publication-quality figures from the existing outputs",
            "Freeze the exact final tables and figure order for the manuscript",
            "Write the Results and Discussion around the canonical final pipeline",
            "State limitations carefully and explicitly",
            "Ideally plan or begin experimental validation to strengthen the paper",
        ],
    )

    doc.add_paragraph(style="Heading 1").add_run("Key Project Files")
    file_par = doc.add_paragraph()
    file_par.add_run("Canonical project summary: ")
    add_hyperlink(file_par, "file://" + str((OUTPUT_DIR / "final_pipeline_summary.md").resolve()), "outputs/final_pipeline_summary.md")
    for path in [
        OUTPUT_DIR / "lead_selection_summary.md",
        OUTPUT_DIR / "final_leads.csv",
        OUTPUT_DIR / "backup_leads.csv",
        OUTPUT_DIR / "enumeration_library_10_methods.csv",
        FINAL_MODEL_DIR / "final_model_report.md",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        add_hyperlink(p, "file://" + str(path.resolve()), str(path.relative_to(ROOT)))

    doc.add_section(WD_SECTION.NEW_PAGE)
    appendix = doc.add_paragraph(style="Heading 1")
    appendix.add_run("Appendix: Internal Text Summaries")
    doc.add_paragraph("The following brief snippets capture the current canonical pipeline and final-model summaries from the project files.")
    doc.add_paragraph("Final pipeline summary", style="Heading 2")
    doc.add_paragraph(pipeline_summary[:2500])
    doc.add_paragraph("Final lead-selection summary", style="Heading 2")
    doc.add_paragraph(lead_summary[:3000])
    doc.add_paragraph("Final model summary", style="Heading 2")
    doc.add_paragraph(model_report[:2500])

    doc.save(REPORT_PATH)
    print(REPORT_PATH)


if __name__ == "__main__":
    build_report()
